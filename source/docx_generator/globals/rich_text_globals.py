import hashlib
import json
import logging
from json import JSONDecodeError
from logging import Logger
from typing import Any, Dict, Union

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docxtpl import Subdoc, DocxTemplate

from docx_generator.adapters.file_adapter import recover_file_path_from_uuid
from docx_generator.globals.utils import retrieve_remote_file, resize_image, get_available_paragraph_alignments


def rendering_decorator(logger: Logger, log_identifier: str, log_start_message: str, log_error_message: str):
    def decorate(func):
        def call():
            logger.debug(f"{log_identifier} [+] {log_start_message}")

            try:
                return func()
            except Exception as e:
                logger.error(f"{log_identifier} - {log_error_message}")
                logger.debug(f"{log_identifier} - Rendering error: {type(e).__name__} {e.__str__()}")
                return None

        return call

    return decorate


def _process_numbered_list_numbering(doc: Document(), par: Paragraph, prev: Paragraph = None, level=None, num=True):
    # Taken from https://github.com/python-openxml/python-docx/issues/25
    """
    Makes a paragraph into a list item with a specific level and
    optional restart.

    An attempt will be made to retrieve an abstract numbering style that
    corresponds to the style of the paragraph. If that is not possible,
    the default numbering or bullet style will be used based on the
    ``num`` parameter.

    :param doc : docx.document.Document
        The document to add the list into.
    :param par : docx.paragraph.Paragraph
        The paragraph to turn into a list item.
    :param prev : docx.paragraph.Paragraph or None
        The previous paragraph in the list. If specified, the numbering
        and styles will be taken as a continuation of this paragraph.
        If omitted, a new numbering scheme will be started.
    :param level : int or None
        The level of the paragraph within the outline. If ``prev`` is
        set, defaults to the same level as in ``prev``. Otherwise,
        defaults to zero.
    :param num : bool
        If ``prev`` is :py:obj:`None` and the style of the paragraph
        does not correspond to an existing numbering style, this will
        determine whether or not the list will be numbered or bulleted.
        The result is not guaranteed, but is fairly safe for most Word
        templates.
    """
    xpath_options = {
        True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
        False: {'single': '', 'level': level},
    }

    def style_xpath(prefer_single=True):
        """
        The style comes from the outer-scope variable ``par.style.name``.
        """
        style = par.style.style_id
        return (
            'w:abstractNum['
            '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
            ']/@w:abstractNumId'
        ).format(style=style, **xpath_options[prefer_single])

    def type_xpath(prefer_single=True):
        """
        The type is from the outer-scope variable ``num``.
        """
        list_type = 'decimal' if num else 'bullet'
        return (
            'w:abstractNum['
            '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
            ']/@w:abstractNumId'
        ).format(type=list_type, **xpath_options[prefer_single])

    def get_abstract_id():
        """
        Select as follows:

            1. Match single-level by style (get min ID)
            2. Match exact style and level (get min ID)
            3. Match single-level decimal/bullet types (get min ID)
            4. Match decimal/bullet in requested level (get min ID)
            3. 0
        """
        for fn in (style_xpath, type_xpath):
            for prefer_single in (True, False):
                xpath = fn(prefer_single)
                ids = numbering.xpath(xpath)
                if ids:
                    return min(int(x) for x in ids)
        return 0

    if (prev is None or
            prev._p.pPr is None or
            prev._p.pPr.numPr is None or
            prev._p.pPr.numPr.numId is None):
        if level is None:
            level = 0
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        # Compute the abstract ID first by style, then by num
        anum = get_abstract_id()
        # Set the concrete numbering based on the abstract numbering ID
        num = numbering.add_num(anum)
        # Make sure to override the abstract continuation property
        num.add_lvlOverride(ilvl=level).add_startOverride(1)
        # Extract the newly-allocated concrete numbering ID
        num = num.numId
    else:
        if level is None:
            level = prev._p.pPr.numPr.ilvl.val
        # Get the previous concrete numbering ID
        num = prev._p.pPr.numPr.numId.val
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level


def _process_paragraph_alignment(paragraph: Paragraph, node_data: Dict):
    child_element_alignment = node_data.get('align', 'left').upper()
    if child_element_alignment in get_available_paragraph_alignments():
        paragraph.alignment = getattr(WD_PARAGRAPH_ALIGNMENT, child_element_alignment)


class RichTextGlobals(object):
    def __init__(
            self,
            template: DocxTemplate,
            base_path: str,
            output_path: str,
            image_directory_path: str,
            style_mapper: Dict[str, str],
            allow_external_download: bool,
            proxy_settings: Dict[str, str]
    ):
        self._template = template

        self._base_path = base_path
        self._output_path = output_path
        self._image_directory_path = image_directory_path

        self._style_mapper = style_mapper
        self._logger = logging.getLogger(__name__)
        self._allow_external_download = allow_external_download
        self._proxy_settings = proxy_settings

    def richtext_to_docx(self, richtext: str) -> Subdoc:
        """
        Converts "RichText" JSON-string data to docx. The data follow the SlateJS WYSIWYG editor data "format" (https://docs.slatejs.org/ - https://www.slatejs.org/examples/richtext).

        The JSON structure is composed of a list of "nodes" having a "type" and several potential "children".
        The generator tries to "translate" these nodes into Docx and tries to set a style that is defined in a dictionary under an eponymous key.
        This dictionary mapping node types to Docx styles is passed as a parameter to the `generate_docx` method.
        The default JSON structure of such node is as follow:

        {"type": "__node_type__", "children": [__some_children__, ...]}

        ---

        **Paragraph Node**

        By default, a node is processed as a paragraph.
        (However it exists several special cases where it does not apply, described underneath)
        The standard structure for a paragraph node is as follow:

        {"type": "__paragraph_node_type__", "children": [{"text": "Same example text"}, ...]}

        The "type" value is used to assign the correct Docx style.
        The "children" value contains a list of all the pieces of text contained in the paragraph.
        It is a list because each piece of text can have different properties: bold, italic, ...
        The description of all the available properties for a text element will be made in a dedicated section.

        Optional properties:

        * align: string - available values: left, center, right, justify, distribute, justify_med, justify_hi, justify_low, thai_justify

        ---

        **Specific Nodes**

        * `image-uuid`

        Include a local image identified by an UUID

        {"type": "image-uuid", "image_uuid": "5bacc2bc-5b90-4c47-93d7-d9291911c4b3"}

        * `image`

        Include a remote image from an address

        {"type": "image", "image_path": "https://]my_image[/]address.jpg"}

        * `caption`

        Add a caption that is going to be recognized as a member of the table of illustrations by Word. The numbering is directly managed by Word.

        {"type": "caption", "text": "Caption example"}

        Optional properties:
            * align: string - available values: left, center, right, justify, distribute, justify_med, justify_hi, justify_low, thai_justify

        * `bulleted-list`

        Add a bulleted list.
        (At the moment rendering process only allows list with 1 level)

        {"type": "bulleted-list", "children": [{"type": "list-item", "children": [{"text": "Item1"}]}, {"type": "list-item", "children": [{"text": "Item2"}]}]}

        * `numbered-list`

        Add a numbered list
        (At the moment rendering process only allows list with 1 level)

        {"type": "numbered-list", "children": [{"type": "list-item", "children": [{"text": "AAA"}]}, {"type": "list-item", "children": [{"text": "BBB"}]}]}

        * `table`

        Add a table.
        Children from a `table` node type must be `table-row` nodes.
        Children from a `table-row` node type must be `table-cell` nodes.

        {type": "table", "children": [{"type": "table-row","children": [{"type": "table-cell","children": [{"type": "paragraph", "children": [{"text": "Same example text"}, ...]}, ...]}

        ---

        **Text Element**

        A text element is used to write some text with specific properties in the docx.
        A default text element looks as follow:

        {"text": "Same example text"}

        Optional properties:

        * `bold`: boolean
        * `italic`: boolean
        * `underline`: boolean
        * `strike`: boolean
        * `code`: boolean - changes the font to `Courier New` and highlights the text in light-gray

        ---

        :param richtext: str
        """
        # What will happen if very large data ?
        logging_identifier = hashlib.sha1(str.encode(richtext)).hexdigest()
        self._logger.info(f"{logging_identifier} - Starting RichText rendering")

        sub_document_template_element = self._template.new_subdoc()
        sub_document_docx_element = sub_document_template_element.subdocx  # type: Document

        last_section = sub_document_docx_element.sections[-1]
        page_width = last_section.page_width - last_section.left_margin - last_section.right_margin

        try:
            json_node_data = json.loads(richtext)
        except JSONDecodeError:
            self._logger.error(f"{logging_identifier} - An error occurred during loading data into JSON")
            json_node_data = [
                {
                    "type": 'paragraph',
                    "children": [{"text": 'An error occurred during JSON parsing'}],
                }
            ]

        def _process_caption_node():
            # caption type
            paragraph = sub_document_docx_element.add_paragraph('Figure ')

            # numbering field
            run = paragraph.add_run()

            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar)

            instrText = OxmlElement('w:instrText')
            instrText.text = f' SEQ Figure \\* ARABIC'
            run._r.append(instrText)

            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar)

            # caption text
            paragraph.add_run(': ')

            return paragraph

        def _process_child(child_element: Dict[str, Any], parent_element: Union[Paragraph, _Cell] = None, forced_style: str = None):
            node_type = child_element.get('type')

            if node_type is not None:
                node_children = child_element.get('children', [{"text": "An error occurred while parsing children"}])
                if node_type == 'image-uuid':

                    @rendering_decorator(self._logger, logging_identifier, 'Rendering image from uuid', 'An error occurred during image rendering from uuid')
                    def _render_image_from_uuid():
                        new_image_paragraph = sub_document_docx_element.add_paragraph()
                        new_image_run = new_image_paragraph.add_run()
                        image_element = new_image_run.add_picture(recover_file_path_from_uuid(self._logger, 'Picture', self._base_path, child_element.get('image_uuid')))
                        if image_element.width > page_width:
                            resize_image(image_element, page_width)

                        return new_image_paragraph

                    return _render_image_from_uuid()

                elif node_type == 'image':

                    @rendering_decorator(self._logger, logging_identifier, 'Rendering image_from_path', 'An error occurred during image processing')
                    def _render_image():
                        local_image = retrieve_remote_file(
                            child_element.get('image_path'),
                            self._base_path,
                            self._image_directory_path,
                            self._logger,
                            self._proxy_settings
                        )

                        new_image_paragraph = sub_document_docx_element.add_paragraph()  # type: Paragraph
                        new_image_run = new_image_paragraph.add_run()
                        image_element = new_image_run.add_picture(local_image)
                        if image_element.width > page_width:
                            resize_image(image_element, page_width)

                        return new_image_paragraph

                    if self._allow_external_download:
                        return _render_image()
                    else:
                        self._logger.info('Cannot download image from outside as it has not been enabled')
                        return None

                elif node_type == 'caption':

                    @rendering_decorator(self._logger, logging_identifier, 'Rendering caption', 'An error occurred during caption rendering')
                    def _render_caption():
                        new_caption_paragraph = _process_caption_node()
                        _process_paragraph_alignment(new_caption_paragraph, child_element)
                        for node_child in node_children:
                            _process_text(new_caption_paragraph, node_child)
                        return new_caption_paragraph

                    return _render_caption()

                elif node_type == 'numbered-list':

                    @rendering_decorator(self._logger, logging_identifier, 'Rendering numbered list', 'An error occurred during numbered list rendering')
                    def _render_numbered_list():
                        new_forced_style = self._style_mapper.get(node_type)
                        previous_child = None
                        for node_child in node_children:
                            new_child = _process_child(node_child, forced_style=new_forced_style, parent_element=parent_element)
                            _process_numbered_list_numbering(sub_document_docx_element, new_child, previous_child, num=True)
                            previous_child = new_child
                        return None

                    return _render_numbered_list()

                elif node_type == 'bulleted-list':

                    @rendering_decorator(self._logger, logging_identifier, 'Rendering bulleted list', 'An error occurred during bulleted list rendering')
                    def _render_bulleted_list():
                        new_forced_style = self._style_mapper.get(node_type)
                        for node_child in node_children:
                            _process_child(node_child, forced_style=new_forced_style, parent_element=parent_element)
                        return None

                    return _render_bulleted_list()

                elif node_type == 'table':

                    @rendering_decorator(self._logger, logging_identifier, 'Rendering table', 'An error occurred during table rendering')
                    def _render_table():
                        table_rows = [row for row in node_children if row.get('type') == 'table-row']
                        nb_table_rows = len(table_rows)
                        nb_table_cells_max = max([len([cell for cell in row.get('children', []) if cell.get('type') == 'table-cell']) for row in table_rows])

                        new_table = sub_document_docx_element.add_table(nb_table_rows, nb_table_cells_max)
                        for index_row, row_child in enumerate(table_rows):
                            self._logger.debug(f"{logging_identifier} \t[+] Processing table row")
                            for index_cell, cell_element in enumerate([cell for cell in row_child.get('children') if cell.get('type') == 'table-cell']):
                                self._logger.debug(f"{logging_identifier} \t\t[+] Processing table cell")
                                considered_cell = new_table.cell(index_row, index_cell)  # type: _Cell

                                cell_children = cell_element.get('children', [])
                                if len(cell_children) > 0:
                                    # By default an empty paragraph is set in each cell. Removing it otherwise an empty line is displayed before the actual content
                                    considered_cell._element.remove(considered_cell.paragraphs[0]._element)
                                    for cell_child in cell_children:
                                        _process_child(cell_child, parent_element=considered_cell)

                        return new_table

                    return _render_table()

                else:

                    @rendering_decorator(self._logger, logging_identifier, f'Rendering {node_type}', f'An error occurred during {node_type} rendering')
                    def _render_node():
                        if parent_element:
                            new_paragraph = parent_element.add_paragraph()
                        else:
                            new_paragraph = sub_document_docx_element.add_paragraph()

                        new_paragraph.style = forced_style if forced_style is not None else self._style_mapper.get(node_type)

                        _process_paragraph_alignment(new_paragraph, child_element)

                        new_forced_style = None

                        for index_child, paragraph_child in enumerate(node_children):
                            if paragraph_child.get('text') is not None:
                                _process_text(new_paragraph, paragraph_child)
                            else:
                                _process_child(paragraph_child, forced_style=new_forced_style)

                        return new_paragraph

                    return _render_node()

        def _process_text(parent_element, text_object: Dict[str, Any]):
            self._logger.debug(f"{logging_identifier} |___ [+] Rendering text value")
            new_run = parent_element.add_run(text_object.get('text', ''))

            font_element = new_run.font
            if text_object.get('bold', False) is True:
                font_element.bold = True
            if text_object.get('italic', False) is True:
                font_element.italic = True
            if text_object.get('underline', False) is True:
                font_element.underline = True
            if text_object.get('strike', False) is True:
                font_element.strike = True
            if text_object.get('code', False) is True:
                # The font should be customizable by the user
                font_element.name = "Courier New"
                font_element.highlight_color = WD_COLOR_INDEX.GRAY_25

            return new_run

        for index, node_element in enumerate(json_node_data):
            _process_child(node_element, forced_style=None)

        return sub_document_template_element
