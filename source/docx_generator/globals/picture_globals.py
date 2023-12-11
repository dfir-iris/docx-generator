#!/usr/bin/env python3
#
#  docx-generator Source Code
#  Copyright (C) 2021 - Airbus CyberSecurity (SAS)
#  ir@cyberactionlab.net
#
#  This program is free software; you can redistribute it and/or
#  modify it under the terms of the GNU Lesser General Public
#  License as published by the Free Software Foundation; either
#  version 3 of the License, or (at your option) any later version.
#
#  This program is distributed in the hope that it will be useful,
#  but WITHOUT ANY WARRANTY; without even the implied warranty of
#  MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the GNU
#  Lesser General Public License for more details.
#
#  You should have received a copy of the GNU Lesser General Public License
#  along with this program; if not, write to the Free Software Foundation,
#  Inc., 51 Franklin Street, Fifth Floor, Boston, MA  02110-1301, USA.

import logging
import os
import re

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docxtpl import DocxTemplate, Subdoc

from docx_generator.adapters.file_adapter import recover_file_path_from_uuid
from docx_generator.exceptions.rendering_error import RenderingError
from docx_generator.globals.utils import retrieve_remote_file, resize_image, get_available_paragraph_alignments


class PictureGlobals(object):
    def __init__(self, template: DocxTemplate, base_path: str, image_directory_path: str):
        self._template = template
        self._base_path = base_path

        self._image_directory_path = image_directory_path

        self._logger = logging.getLogger(__name__)

    def set_template(self, template: DocxTemplate):
        self._template = template

    def set_base_path(self, base_path: str):
        self._base_path = base_path

    def set_image_directory_path(self, output_path: str):
        try:
            if not os.path.isdir(output_path):
                os.mkdir(output_path)

            return output_path
        except Exception as e:
            self._logger.critical("Impossible to create temporary image directory")
            self._logger.debug(f"Image directory error: {e.__str__()}")
            return None

    def _process_image(self, position, image_filename: str) -> Subdoc:
        sub_document = self._template.new_subdoc()

        last_section = sub_document.sections[-1]
        page_width = last_section.page_width - last_section.left_margin - last_section.right_margin

        try:
            picture = sub_document.add_picture(image_filename)
        except Exception as e:
            self._logger.debug('Error while adding image {}: {}'.format(image_filename, e.__str__()))
            self._logger.debug('There is a problem sometimes with JPEG-Files and EXIF-Headers, try a PNG instead')
            raise RenderingError(self._logger, 'Image could not be added (try PNG instead of JPEG): {}'.format(image_filename))

        # Scale picture to page dimension if width is bigger than page width
        if picture.width > page_width:
            resize_image(picture, page_width)
            self._logger.info(f"Image resized - {image_filename}")

        if position in get_available_paragraph_alignments():
            last_paragraph = sub_document.paragraphs[-1]
            last_paragraph.alignment = getattr(WD_PARAGRAPH_ALIGNMENT, position)

        self._logger.debug('Image added: {} {}'.format(position, image_filename))

        return sub_document

    def add_picture(self, image_path: str, position: str = 'CENTER') -> Subdoc:
        """
        Adds picture to document from local path.

        :param image_path: str
            Full path to picture.
        :param position: str
            Position value used to position value in the document.
            Available values:
            ['LEFT', 'CENTER', 'RIGHT', 'JUSTIFY', 'DISTRIBUTE', 'JUSTIFY_MED', 'JUSTIFY_HI', 'JUSTIFY_LOW', 'THAI_JUSTIFY']
            (Default value: CENTER)

        :return: docxtpl.Subdoc
        """

        try:
            image_path = retrieve_remote_file(image_path, self._base_path, self._image_directory_path, self._logger)
        except Exception:
            self._logger.error(f'Skipping {image_path} due to error')
            return self._template.new_subdoc()

        return self._process_local(image_path, position)

    def _process_local(self, image_path: str, position: str = 'CENTER') -> Subdoc:
        """
        Process the image as a locally stored file.
        """
        incorrect_path_pattern = r'\.\.'
        if len(re.findall(incorrect_path_pattern, image_path)) > 0:
            raise RenderingError(self._logger, 'Invalid filename provided')

        if not os.path.isfile(image_path):
            raise RenderingError(self._logger, 'The path provided is not a correct file', 'The path provided is not a correct file: {}'.format(image_path))

        try:
            return_value = self._process_image(position, image_path)
            return return_value
        except Exception as e:
            raise RenderingError(self._logger, e.__str__())

    def add_picture_from_uuid(self, uuid: str, position: str = 'CENTER') -> Subdoc:
        """
        Adds picture to document from special file structure. Images must be stored in a folder being named with a uuid identifying the picture. This folder must be stored directly under the bas path.
        :
        uuid:   str
            Uuid of the picture.
            example: 466cf6e1-569d-4239-ae34-9a4d9b52fd5c
        :position:   str
            Position value used to position value in the document.
            Available values:
            ['LEFT', 'CENTER', 'RIGHT', 'JUSTIFY', 'DISTRIBUTE', 'JUSTIFY_MED', 'JUSTIFY_HI', 'JUSTIFY_LOW', 'THAI_JUSTIFY']
            (Default value: CENTER)

        :return: docxtpl.Subdoc

        """
        picture_file_path = recover_file_path_from_uuid(self._logger, 'Picture', self._base_path, uuid)

        return self._process_image(position, picture_file_path)
