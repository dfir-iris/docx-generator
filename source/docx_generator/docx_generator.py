#!/usr/bin/env python3
#
#  docx-generator Source Code
#  Copyright (C) 2021 - Airbus CyberSecurity (SAS)
#  ir@cyberactionlab.net
#
#  This program is free software; you can redistribute it and/or
#  modify it under the terms of the GNU Lesser General Public
#  License as published by the Free Software Foundation; either
#  version 3 of the License, or (at your option) any later version.
#
#  This program is distributed in the hope that it will be useful,
#  but WITHOUT ANY WARRANTY; without even the implied warranty of
#  MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the GNU
#  Lesser General Public License for more details.
#
#  You should have received a copy of the GNU Lesser General Public License
#  along with this program; if not, write to the Free Software Foundation,
#  Inc., 51 Franklin Street, Fifth Floor, Boston, MA  02110-1301, USA.

import logging
import os
import re
from typing import Dict, List

import errno
from docx.styles.styles import Styles
from docxtpl import DocxTemplate
from jinja2 import Environment

from docx_generator.adapters.docx.style_adapter import RenderStylesCollection, get_document_render_styles
from docx_generator.adapters.mistletoe.DocxRenderer import DocxRenderer
from docx_generator.exceptions.rendering_error import RenderingError
from docx_generator.filters.filters import Filters
from docx_generator.globals.globals import Globals
from docx_generator.globals.picture_globals import PictureGlobals


def _sanitize_path(path: str) -> str:
    sanitized_path = path.strip('./').strip(',').strip(' ')
    return os.path.normpath(sanitized_path)


class DocxGenerator(object):
    def __init__(
        self,
        logger_mode: str = 'INFO',
        max_recursive_render_depth: int = 5,
        image_handler: PictureGlobals = None,
        app_logger: logging = None,
        allow_external_download: bool = False,
        proxy_settings: Dict[str, str] = None
    ):
        """
        :type logger_mode: Enum(DEBUG, INFO, WARNING, CRITICAL)
        :param logger_mode: Sets the logger level
        :type max_recursive_render_depth: integer
        :param max_recursive_render_depth:
            It is possible to embed Jinja2 variables inside data sent to the generator. After the document generation, the generator look for remaining variables in the generated document, and does another generation with the same data if it finds at least one. This parameter determines how many time the generator will perform this process
        :type image_handler:
        :param image_handler:
        :type app_logger: Logger
        :param app_logger:
            If the generator is used in the context of another application, it allows to pass the application's logger so that the generator writes its logs to it.
        :type allow_external_download: boolean
        :param allow_external_download:
            Used to determine if the generator can query elements externally, for example images directly from the internet.
        :type proxy_settings: Dict
        :param proxy_settings:
            Used for setting HTTP and HTTPS proxy. Value should look like {'http': '__HTTP_PROXY_ADDRESS__', 'https': '__HTTPS_PROXY_ADDRESS__'}
        """
        if app_logger is None:
            logging.basicConfig(
                format='%(asctime)s :: %(levelname)s :: %(name)s :: %(message)s',
                level=getattr(logging, logger_mode, logging.INFO)
            )

            self._logger = logging.getLogger(__name__)
        else:
            self._logger = app_logger

        self._max_recursive_render_depth = max_recursive_render_depth
        self._image_handler = image_handler
        self._allow_external_download = allow_external_download
        self._proxy_settings = proxy_settings if proxy_settings is not None else {}

    def _process_template_path(self, base_path: str, template_path: str) -> str:
        template_path = _sanitize_path(template_path)
        full_template_path = os.path.join(base_path, template_path)
        if not os.path.isfile(full_template_path):
            raise RenderingError(self._logger, 'Generator can not find template.', 'Generator can not find template: {}'.format(full_template_path))
        else:
            self._logger.info('Template located: {}'.format(full_template_path))

        return full_template_path

    def _process_output_path(self, base_path: str, output_path: str) -> str:
        output_path = _sanitize_path(output_path)
        full_output_path = os.path.join(base_path, output_path)
        full_output_dir = os.path.dirname(full_output_path)
        if not os.path.isdir(full_output_dir):
            raise RenderingError(self._logger, 'Generator can not find output directory.', 'Generator can not find output directory to create {}'.format(full_output_path))
        else:
            self._logger.info('Output directory located: {}'.format(full_output_path))
        return full_output_path

    def _process_image_directory_path(self, base_path: str, image_directory_path: str) -> str:
        if image_directory_path is None:
            image_directory_path = os.path.join('tmp', 'images')

        full_image_directory_path = os.path.join(base_path, image_directory_path)

        try:
            if not os.path.isdir(full_image_directory_path):
                os.makedirs(full_image_directory_path)

            return full_image_directory_path
        except OSError as exc:
            if exc.errno == errno.EEXIST and os.path.isdir(full_image_directory_path):
                return full_image_directory_path
        except Exception as e:
            self._logger.critical("Impossible to create temporary image directory")
            self._logger.debug(f"Image directory error: {e.__str__()}")
            raise RenderingError(self._logger, 'Image directory path passed to the generator is not valid')

    def _set_jinja2_custom_environment(self, base_path: str, output_path: str, image_directory_path: str, template: DocxTemplate, style_mapper: Dict[str, str], jinja2_environment: Environment, renderer: DocxRenderer, template_styles: RenderStylesCollection) -> None:
        jinja2_custom_filters = Filters(renderer, template_styles, jinja2_environment)
        jinja2_custom_globals = Globals(
            base_path,
            output_path,
            image_directory_path,
            template, style_mapper,
            jinja2_environment,
            self._allow_external_download,
            self._proxy_settings
        )

        jinja2_custom_filters.set_available_filters()
        jinja2_custom_globals.set_available_globals()

    def _recursive_rendering(self, base_path: str, template_path: str, style_mapper: Dict[str, str], data: Dict, output_path: str, image_directory_path: str, render_level: int):
        render_level += 1
        self._logger.info('Start rendering for level {}'.format(render_level))

        loaded_template = DocxTemplate(template_path)
        template_styles = get_document_render_styles(template_path)

        docx_renderer = DocxRenderer(loaded_template, self._image_handler)

        jinja_custom_environment = Environment()

        self._set_jinja2_custom_environment(base_path, output_path, image_directory_path, loaded_template, style_mapper, jinja_custom_environment, docx_renderer, template_styles)

        try:
            loaded_template.render(data, jinja_env=jinja_custom_environment, autoescape=True)
        except Exception as e:
            error_message = '{} ({})'.format(str(e), os.path.basename(template_path))
            raise RenderingError(self._logger, error_message)

        is_variable_found = False
        variable_regex = "{{.+}}|{%.+%}"

        for paragraph in loaded_template.paragraphs:
            if re.search(variable_regex, paragraph.text) is not None:
                is_variable_found = True
                break

        for table in loaded_template.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if re.search(variable_regex, paragraph.text) is not None:
                            is_variable_found = True
                            break

        loaded_template.save(output_path)
        self._logger.info('Document generated for level {}'.format(render_level))

        if is_variable_found and render_level <= self._max_recursive_render_depth:
            self._logger.info('Variable found in generated document. Restarting rendering process ...')
            self._recursive_rendering('', output_path, style_mapper, data, output_path, image_directory_path, render_level)

        if render_level > self._max_recursive_render_depth:
            self._logger.info('Rendering depth level exceeded, leaving render loop')

        self._logger.info('Rendering process completed !')

    """
        template_path and absolute_path must be relative to base_path
    """

    def generate_docx(self, base_path: str, template_path: str, data: Dict, output_path: str, style_mapper: Dict[str, str] = None, image_directory_path: str = None):
        processed_base_path = os.path.abspath(base_path)
        full_template_path = self._process_template_path(processed_base_path, template_path)
        full_output_path = self._process_output_path(processed_base_path, output_path)
        full_image_directory_path = self._process_image_directory_path(processed_base_path, image_directory_path)

        style_mapper = {} if style_mapper is None else style_mapper

        if self._image_handler is not None:
            self._image_handler.set_base_path(processed_base_path)
            self._image_handler.set_image_directory_path(os.path.join(os.path.dirname(full_output_path), "images"))

        self._logger.info('Starting new report generation. Base path: {}. Template path: {}. Output path'.format(processed_base_path, full_template_path, full_output_path))
        self._recursive_rendering(processed_base_path, full_template_path, style_mapper, data, full_output_path, full_image_directory_path, 0)

    def get_available_styles(self, base_path: str, template_path: str) -> List[Dict[str, str]]:
        processed_base_path = os.path.abspath(base_path)
        full_template_path = self._process_template_path(processed_base_path, template_path)

        styles = DocxTemplate(full_template_path).get_docx().styles  # type: Styles

        return [{"name": style.name, "id": style.style_id} for style in styles]
