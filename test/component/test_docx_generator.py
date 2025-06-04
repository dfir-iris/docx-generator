#!/usr/bin/env python3
#
#  docx-generator Source Code
#  Copyright (C) 2021 - Airbus CyberSecurity (SAS)
#  ir@cyberactionlab.net
#
#  This program is free software; you can redistribute it and/or
#  modify it under the terms of the GNU Lesser General Public
#  License as published by the Free Software Foundation; either
#  version 3 of the License, or (at your option) any later version.
#
#  This program is distributed in the hope that it will be useful,
#  but WITHOUT ANY WARRANTY; without even the implied warranty of
#  MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the GNU
#  Lesser General Public License for more details.
#
#  You should have received a copy of the GNU Lesser General Public License
#  along with this program; if not, write to the Free Software Foundation,
#  Inc., 51 Franklin Street, Fifth Floor, Boston, MA  02110-1301, USA.

import os
import re
from pathlib import Path
from unittest import TestCase

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE

from docx_generator.docx_generator import DocxGenerator
from docx_generator.exceptions.rendering_error import RenderingError


class TestDocxGenerator(TestCase):
    def setUp(self) -> None:
        self._base_path = os.path.join(os.getcwd(), 'test/component')
        self._template_path = 'templates'
        self._results_path = 'results'

        self._output_filenames = {
            'basic_template_result': 'basic_template_result.docx',
            'date_filter_template_result': 'date_filter_template_result.docx',
            'image_filter_template_result': 'image_filter_template_result.docx',
            'image_from_uuid_filter_template_result': 'image_from_uuid_filter_template_result.docx',
            'markdown_filter_template_result': 'markdown_filter_template_result.docx',
            'specific_markdown_filter_template_result': 'specific_markdown_filter_template_result.docx',
            'subdoc_filter_template_result': 'subdoc_filter_template_result.docx',
            'recursive_render_result': 'recursive_render_result.docx',
            'recursive_max_depth_render_result': 'recursive_max_depth_render_result.docx',
            'hyperlink_global_result': 'hyperlink_global_result.docx',
            'non_existent_global_result': 'non_existent_global_result.docx',
            'non_existent_filter_result': 'non_existent_filter_result.docx',
            'unclosed_jinja_control_tag_result': 'unclosed_jinja_control_tag_result.docx'
        }

        Path(self._base_path, self._results_path).mkdir(exist_ok=True)
        self._subject = DocxGenerator(logger_mode='DEBUG')

    def tearDown(self) -> None:
        return None
        for filename in self._output_filenames.values():
            try:
                os.remove(os.path.join(self._results_path, filename))
            except FileNotFoundError:
                pass

    def test_should_generate_docx_from_basic_template(self):
        text_to_add = 'Report Name'
        data = {
            'name': text_to_add
        }

        self._subject.generate_docx(
            self._base_path,
            os.path.join(self._template_path, 'basic_template.docx'),
            data,
            os.path.join(self._results_path, self._output_filenames['basic_template_result'])
        )

        has_been_found = False
        document = Document(os.path.join(self._base_path, self._results_path, self._output_filenames['basic_template_result']))
        for paragraph in document.paragraphs:
            if re.search(text_to_add, paragraph.text) is not None:
                has_been_found = True
                break

        self.assertTrue(has_been_found, 'Text passed as data is not found in the generated document.')

    def test_should_generate_docx_from_template_with_timestamp_to_date_filter(self):
        data = {
            'date': '1589480671562'
        }

        converted_date = '14/05/2020'

        self._subject.generate_docx(
            self._base_path,
            os.path.join(self._template_path, 'date_filter_template.docx'),
            data,
            os.path.join(self._results_path, self._output_filenames['date_filter_template_result'])
        )

        has_been_found = False
        document = Document(os.path.join(self._base_path, self._results_path, self._output_filenames['date_filter_template_result']))
        for paragraph in document.paragraphs:
            if re.search(converted_date, paragraph.text) is not None:
                has_been_found = True
                break

        self.assertTrue(has_been_found, 'Converted date is not found in the generated document.')

    def test_should_generate_docx_from_template_with_hyperlink_global(self):
        data = {
            'hyperlink_caption': 'google',
            'hyperlink_url': 'https://www.google.fr',

            'mail_caption': 'alain.dupont@orange.fr',
            'mail_url': 'mailto:alain.dupont@orange.fr'
        }

        self._subject.generate_docx(
            self._base_path,
            os.path.join(self._template_path, 'hyperlink_global_template.docx'),
            data,
            os.path.join(self._results_path, self._output_filenames['hyperlink_global_result'])
        )

        # TESTS
        generated_document = Document(os.path.join(self._base_path, self._results_path, self._output_filenames['hyperlink_global_result']))

        found_hyperlinks = []

        relationships = generated_document.part.rels
        for rel in relationships:
            if relationships[rel].reltype == RELATIONSHIP_TYPE.HYPERLINK:
                found_hyperlinks.append(relationships[rel]._target)

        self.assertEqual(2, len(found_hyperlinks))
        self.assertTrue(data['hyperlink_url'] in found_hyperlinks)
        self.assertTrue(data['mail_url'] in found_hyperlinks)

    def test_should_generate_docx_from_template_with_image_global(self):
        data = {
            'image1': os.path.abspath(os.path.join(self._base_path, './images/test_image.jpg')),
            'image2': os.path.abspath(os.path.join(self._base_path, './images/test_image_small.jpg'))
        }

        self._subject.generate_docx(
            self._base_path,
            os.path.join(self._template_path, 'image_filter_template.docx'),
            data,
            os.path.join(self._results_path, self._output_filenames['image_filter_template_result'])
        )

    def test_should_generate_docx_from_template_with_image_from_uuid_global(self):
        data = {
            'uuid': '5bacc2bc-5b90-4c47-93d7-d9291911c4b3',
        }

        self._subject.generate_docx(
            self._base_path,
            os.path.join(self._template_path, 'image_from_uuid_filter_template.docx'),
            data,
            os.path.join(self._results_path, self._output_filenames['image_from_uuid_filter_template_result'])
        )

    def test_should_raise_error_if_uuid_folder_contains_multiple_files(self):
        data = {
            'uuid': '2b910dac-6e10-45ae-89cd-b9c304809eb9',
        }

        with self.assertRaises(RenderingError):
            self._subject.generate_docx(
                self._base_path,
                os.path.join(self._template_path, 'image_from_uuid_filter_template.docx'),
                data,
                os.path.join(self._results_path, self._output_filenames['image_from_uuid_filter_template_result'])
            )

    def test_should_generate_docx_from_template_with_markdown_filter(self):
        markdown_text = 'First Paragraph:\n\n**Strong text**\n*New line with italic text* and `code`\n\nNew paragraph\n\n' \
                        '[link to Google!](http://google.com )\n\nUnordered List:\n\n* Item 1\n* Item 2\n\nOrdered List\n\n' \
                        '1. Item A\n2. Item B\n\nA block of code:\n\n```\ntest with some code\nAnother line of code\n__should not be bold__\n' \
                        '```\n\nTable:\n\n**Markdown** | *Less* | Pretty\n--- | --- | ---\n1 | 2 | 3\n\n \n\n' \
                        'Quote:\n\n> Is a quote !\n> Is still a quote !!\n> **Strong quote**\n' \
                        '```\nPOST /aaa/bbb/ccc\nHTTP / 1.1\nHost: toto.toto.toto\nUser - Agent: curl / 7.58.0 \n' \
                        'Accept: * / *\nContent - Type: application / json\nAuthorization: Bearer c66330ee-cd76-4ab--d37d44bef\n' \
                        'Id_token: __token__\nContent - Length: 305\nConnection: close\n\n{"key1": "value1",\n"key2": "value2\n}\n```'

        markdown_text2 = 'toto'

        data = {
            'text_for_paragraph': markdown_text,
            'text_for_code_block': markdown_text2
        }

        self._subject.generate_docx(
            self._base_path,
            os.path.join(self._template_path, 'markdown_filter_template.docx'),
            data,
            os.path.join(self._results_path, self._output_filenames['markdown_filter_template_result'])
        )
    def test_should_not_fail_with_specific_markdown(self):
        markdown_text = '***possibly an error***'

        markdown_text2 = 'toto'

        data = {
            'text_for_paragraph': markdown_text,
            'text_for_code_block': markdown_text2
        }

        self._subject.generate_docx(
            self._base_path,
            os.path.join(self._template_path, 'markdown_filter_template.docx'),
            data,
            os.path.join(self._results_path, self._output_filenames['specific_markdown_filter_template_result'])
        )

    def test_should_generate_docx_from_template_with_subdocument_global(self):
        subdoc_path = os.path.join(self._base_path, self._template_path, 'sub_document_filter_template_part.docx')

        data = {
            'sub_document_path': subdoc_path
        }

        self._subject.generate_docx(
            self._base_path,
            os.path.join(self._template_path, 'sub_document_filter_template.docx'),
            data,
            os.path.join(self._results_path, self._output_filenames['subdoc_filter_template_result'])
        )

    def test_should_generate_docx_with_nested_variables(self):
        subdoc_path = os.path.join(self._base_path, self._template_path, 'sub_document_filter_template_part_with_nested_variable.docx')

        data = {
            'sub_document_path': subdoc_path,
            'nested_variable': 'A nested value included {{nested_variable_level_2}}',
            'nested_variable_level_2': 'with some more one level under.'
        }

        self._subject.generate_docx(
            self._base_path,
            os.path.join(self._template_path, 'sub_document_filter_template.docx'),
            data,
            os.path.join(self._results_path, self._output_filenames['recursive_render_result'])
        )

    def test_should_generate_docx_with_nested_variables_up_to_5_render(self):
        subdoc_path = os.path.join(self._base_path, self._template_path, 'sub_document_filter_template_part_with_nested_variable.docx')

        data = {
            'sub_document_path': subdoc_path,
            'nested_variable': 'A nested value included {{nested_variable_level_2}}',
            'nested_variable_level_2': 'with some more from level 2 {{nested_variable_level_3}}',
            'nested_variable_level_3': 'with some more from level 3 {{nested_variable_level_4}}',
            'nested_variable_level_4': 'with some more from level 4 {{nested_variable_level_5}}',
            'nested_variable_level_5': 'with some more from level 5 {{nested_variable_level_6}}',
            'nested_variable_level_6': 'with some more from level 6',
        }

        self._subject.generate_docx(
            self._base_path,
            os.path.join(self._template_path, 'sub_document_filter_template.docx'),
            data,
            os.path.join(self._results_path, self._output_filenames['recursive_max_depth_render_result'])
        )

    def test_should_raise_rendering_error_if_global_does_not_exist(self):
        data = {'value': 'test values'}

        with self.assertRaises(RenderingError):
            self._subject.generate_docx(
                self._base_path,
                os.path.join(self._template_path, 'non_existent_global_template.docx'),
                data,
                os.path.join(self._results_path, self._output_filenames['non_existent_global_result'])
            )

    def test_should_raise_rendering_error_if_filter_does_not_exist(self):
        data = {'value': 'test values'}

        with self.assertRaises(RenderingError):
            self._subject.generate_docx(
                self._base_path,
                os.path.join(self._template_path, 'non_existent_filter_template.docx'),
                data,
                os.path.join(self._results_path, self._output_filenames['non_existent_filter_result'])
            )

    def test_should_raise_rendering_error_if_there_is_an_error_with_jinja_control_tags(self):
        data = {'value': 'test values'}

        with self.assertRaises(RenderingError):
            self._subject.generate_docx(
                self._base_path,
                os.path.join(self._template_path, 'unclosed_jinja_control_tag_template.docx'),
                data,
                os.path.join(self._results_path, self._output_filenames['unclosed_jinja_control_tag_result'])
            )
